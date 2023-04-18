import docx

# Open the document
doc = docx.Document('Radio4.docx')

# Create a dictionary to store the paragraph text and its count
paragraph_dict = {}

# Loop through the paragraphs in the document and count the repeated ones
for paragraph in doc.paragraphs:
    text = paragraph.text.strip()
    if text not in paragraph_dict:
        paragraph_dict[text] = 1
    else:
        paragraph_dict[text] += 1

# Loop through the paragraphs again and make the repeated ones bold and italic
for paragraph in doc.paragraphs:
    text = paragraph.text.strip()
    if paragraph_dict[text] > 1:
        paragraph.style.font.bold = True
        paragraph.style.font.italic = True
        # Decrement the count of repeated paragraphs to avoid making them bold and italic again
        paragraph_dict[text] -= 1

# Save the document
doc.save('your_updated_docx_file.docx')
