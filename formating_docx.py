import docx

# Open the input file
input_doc = docx.Document('Radio4.docx')

# Create a new document for modified content
output_doc = docx.Document()

# Create a new style for bold and italic font
italic_style = output_doc.styles.add_style('MyBoldItalic', docx.enum.style.WD_STYLE_TYPE.PARAGRAPH)
italic_style.font.italic = True
italic_style.font.bold = True

# Loop through the paragraphs and apply formatting
for i, para in enumerate(input_doc.paragraphs):
    # Determine the formatting to use
    if i % 2 == 0:
        # Even paragraphs are normal font
        font_format = None
    else:
        # Odd paragraphs are bold and italic
        font_format = 'MyBoldItalic'

    # Add the paragraph to the new document with the desired formatting
    output_para = output_doc.add_paragraph(style=font_format)
    output_para.add_run(para.text)

# Save the modified document
output_doc.save('output.docx')
