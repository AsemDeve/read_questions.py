import re
import os
import errno
import Levenshtein
import docx2txt
import csv_operations
import webbrowser
import docx

csv_header = ['question', 'ratio']
doc_or = docx2txt.process('test.docx')
pattern_q = r"#(.*?)#"
choices_pattern = '[a-fA-F][.|-|)|:].*'

output_doc = docx.Document()

# Create a new style for bold and italic font
italic_style = output_doc.styles.add_style('MyBoldItalic', docx.enum.style.WD_STYLE_TYPE.PARAGRAPH)
italic_style.font.italic = True
italic_style.font.bold = True

while True:
    try:
        # ratio = float(input("Enter a ratio between 0 and 1: "))
        ratio = 0.9
        if ratio < 0 or ratio > 1:
            print("Error: Number must be between 0 and 1")
        else:
            break
    except ValueError:
        print("Error: Input must be a number between 0 and 1")

question_text = re.findall(pattern_q, doc_or, flags=re.DOTALL)
my_res = []
count = 0

for i in range(len(question_text)):
    # font_format = None
    for j in range(i + 1, len(question_text)):
        text_ratio = Levenshtein.ratio(question_text[j][3:].lower().strip(), question_text[i][3:].lower().strip())
        if text_ratio >= ratio:
            my_res.append([question_text[i].strip(), round(text_ratio, 4)])
            my_res.append([question_text[j].strip(), ''])
            my_res.append(["***********************", "*****"])
            count += 1
            font_format = 'MyBoldItalic'
            output_para = output_doc.add_paragraph(style=font_format)
            output_para.add_run(question_text[j])

        else:
            font_format = None
            output_para = output_doc.add_paragraph(style=font_format)
            output_para.add_run(question_text[j])

    # output_para = output_doc.add_paragraph(style=font_format)
    # output_para.add_run(question_text[j])

print("numbers of duplicate items : " + str(count))
#
csv_operations.writer('duplicate_in_docx.csv', my_res, csv_header)
cwd = os.getcwd()
csv_file_path = os.path.join(cwd, 'duplicate_in_docx.csv')
output_doc.save('output.docx')
docx_file_path = os.path.join(cwd, 'output.docx')

try:
    webbrowser.open('file://' + csv_file_path)
    webbrowser.open('file://' + docx_file_path)

except FileNotFoundError:
    print("The file could not be found.")
except IOError:
    print("An error occurred while trying to open the file.")
except OSError as e:
    if e.errno == errno.EACCES:
        print("Permission denied. You do not have the necessary permissions to access this file.")
